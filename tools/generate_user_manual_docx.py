"""
生成《通用CAE分析平台软件使用说明书》Word 文档。

说明：
1. 本脚本基于当前仓库中的真实实现整理说明书内容，避免文档与软件功能脱节；
2. 输出文件默认覆盖用户指定的软著说明书路径；
3. 文中的图片先统一生成占位图，后续可在 Word 中替换为真实截图。
"""

from __future__ import annotations

import argparse
import shutil
import textwrap
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PLACEHOLDER_DIR = PROJECT_ROOT / "docs" / "manual_placeholders_v2"
DEFAULT_OUTPUT = Path(r"D:\研究生课外\CAE通用分析平台软件\通用CAE分析平台软件_使用说明书.docx")
DEFAULT_TITLE = "通用CAE分析平台软件"
DEFAULT_VERSION = "V1.0"


def _load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    """加载占位图使用的中文字体。"""

    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simkai.ttf",
    ]
    for font_path in candidates:
        path = Path(font_path)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def _ensure_placeholder(
    key: str,
    title: str,
    hint_lines: Sequence[str],
    accent_rgb: tuple[int, int, int],
) -> Path:
    """生成截图占位图。"""

    PLACEHOLDER_DIR.mkdir(parents=True, exist_ok=True)
    image_path = PLACEHOLDER_DIR / f"{key}.png"
    if image_path.exists():
        return image_path

    width, height = 1600, 900
    background = (245, 247, 250)
    accent_light = tuple(min(255, value + 120) for value in accent_rgb)
    image = Image.new("RGB", (width, height), background)
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((70, 60, width - 70, 180), radius=32, fill=accent_rgb)
    title_font = _load_font(48, bold=True)
    draw.text((110, 92), title, font=title_font, fill=(255, 255, 255))

    draw.rounded_rectangle((70, 230, width - 70, height - 130), radius=32, outline=accent_rgb, width=6, fill=(255, 255, 255))
    for x_pos in range(120, width - 120, 50):
        draw.line((x_pos, 310, x_pos + 18, 310), fill=accent_light, width=4)
        draw.line((x_pos, height - 210, x_pos + 18, height - 210), fill=accent_light, width=4)
    for y_pos in range(310, height - 210, 42):
        draw.line((120, y_pos, 120, y_pos + 16), fill=accent_light, width=4)
        draw.line((width - 120, y_pos, width - 120, y_pos + 16), fill=accent_light, width=4)

    center_title_font = _load_font(58, bold=True)
    center_note_font = _load_font(30, bold=False)
    box_title = "截图占位图"
    center_box = draw.textbbox((0, 0), box_title, font=center_title_font)
    draw.text(((width - (center_box[2] - center_box[0])) / 2, 360), box_title, font=center_title_font, fill=accent_rgb)

    wrapped_lines: list[str] = []
    for line in hint_lines:
        wrapped_lines.extend(textwrap.wrap(line, width=26, break_long_words=False) or [""])
    wrapped_lines.append("后续可直接替换为对应模块的真实运行截图。")

    current_y = 455
    for line in wrapped_lines:
        line_box = draw.textbbox((0, 0), line, font=center_note_font)
        draw.text(((width - (line_box[2] - line_box[0])) / 2, current_y), line, font=center_note_font, fill=(74, 85, 104))
        current_y += 46

    footer_font = _load_font(24, bold=False)
    footer = "本图为自动生成占位图，用于后续替换真实界面截图"
    footer_box = draw.textbbox((0, 0), footer, font=footer_font)
    draw.text(((width - (footer_box[2] - footer_box[0])) / 2, height - 92), footer, font=footer_font, fill=(104, 117, 134))

    image.save(image_path)
    return image_path


def _set_run_font(run, font_name: str = "宋体", size: float = 11, bold: bool | None = None, color: RGBColor | None = None) -> None:
    """统一设置 run 的中西文字体。"""

    run.font.name = font_name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def _set_paragraph_font(paragraph, font_name: str = "宋体", size: float = 11, bold: bool | None = None, color: RGBColor | None = None) -> None:
    """统一设置段落内所有 run 的字体。"""

    if not paragraph.runs:
        run = paragraph.add_run("")
        _set_run_font(run, font_name=font_name, size=size, bold=bold, color=color)
        return
    for run in paragraph.runs:
        _set_run_font(run, font_name=font_name, size=size, bold=bold, color=color)


def _style_document(document: Document) -> None:
    """设置页面和主要样式。"""

    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for style_name, font_name, size, bold in [
        ("Title", "黑体", 22, True),
        ("Heading 1", "黑体", 16, True),
        ("Heading 2", "黑体", 14, True),
        ("Heading 3", "黑体", 12, True),
        ("Subtitle", "宋体", 14, False),
        ("List Bullet", "宋体", 11, False),
        ("List Number", "宋体", 11, False),
    ]:
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_format(paragraph, first_line_cm: float = 0.74, space_before: float = 0, space_after: float = 6) -> None:
    """统一正文段落格式。"""

    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(first_line_cm)
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_paragraph(document: Document, text: str, *, indent: bool = True) -> None:
    """添加普通正文段落。"""

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run, font_name="宋体", size=11)
    _set_paragraph_format(paragraph, first_line_cm=0.74 if indent else 0.0)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    """添加项目符号列表。"""

    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = paragraph.add_run(item)
        _set_run_font(run, font_name="宋体", size=11)


def add_numbered(document: Document, items: Iterable[str]) -> None:
    """添加编号列表。"""

    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = paragraph.add_run(item)
        _set_run_font(run, font_name="宋体", size=11)


def add_table(document: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths_cm: Sequence[float] | None = None) -> None:
    """添加带边框表格。"""

    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for column_index, header in enumerate(headers):
        cell = table.cell(0, column_index)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        _set_run_font(run, font_name="黑体", size=10.5, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths_cm and column_index < len(widths_cm):
            cell.width = Cm(widths_cm[column_index])

    for row_index, row in enumerate(rows, start=1):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            _set_run_font(run, font_name="宋体", size=10.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths_cm and column_index < len(widths_cm):
                cell.width = Cm(widths_cm[column_index])

    document.add_paragraph("")


def add_heading(document: Document, text: str, level: int) -> None:
    """添加标题并统一样式。"""

    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    font_name = "黑体"
    size = 16 if level == 1 else 14 if level == 2 else 12
    _set_paragraph_font(paragraph, font_name=font_name, size=size, bold=True)


def _make_field_run(paragraph, instruction: str) -> None:
    """向段落中写入一个 Word 域。"""

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def add_toc(document: Document) -> None:
    """插入目录域。"""

    add_heading(document, "目录", 1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _make_field_run(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run("说明：若目录页码未自动刷新，请在 Word 中右键目录并选择“更新域”。")
    _set_run_font(run, font_name="宋体", size=10.5, color=RGBColor(90, 101, 115))
    note.paragraph_format.space_after = Pt(12)
    document.add_page_break()


def add_footer_page_number(section) -> None:
    """为节页脚添加页码。"""

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    _set_run_font(run, font_name="宋体", size=10.5)
    _make_field_run(paragraph, "PAGE")
    run = paragraph.add_run(" 页 / 共 ")
    _set_run_font(run, font_name="宋体", size=10.5)
    _make_field_run(paragraph, "NUMPAGES")
    run = paragraph.add_run(" 页")
    _set_run_font(run, font_name="宋体", size=10.5)


def add_cover(document: Document, title: str, version: str) -> None:
    """生成封面页。"""

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(90)
    run = paragraph.add_run(title)
    _set_run_font(run, font_name="黑体", size=24, bold=True, color=RGBColor(31, 41, 55))

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    run = paragraph.add_run("软件使用说明书")
    _set_run_font(run, font_name="宋体", size=18, bold=True, color=RGBColor(37, 99, 235))

    meta_lines = [
        f"软件版本：{version}",
        "文档版本：2026 年 03 月修订版",
        f"编写日期：{datetime.now().strftime('%Y年%m月%d日')}",
        "适用对象：软件著作权申报、发布版交付、日常使用培训",
    ]
    for line in meta_lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(10)
        run = paragraph.add_run(line)
        _set_run_font(run, font_name="宋体", size=13)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(60)
    run = paragraph.add_run("说明：本文档中的界面图片均为占位图，后续可按图题替换为真实截图。")
    _set_run_font(run, font_name="宋体", size=11, color=RGBColor(90, 101, 115))

    document.add_page_break()


def add_figure(
    document: Document,
    *,
    key: str,
    figure_no: str,
    title: str,
    hint_lines: Sequence[str],
    accent_rgb: tuple[int, int, int] = (37, 99, 235),
    width_inches: float = 6.2,
) -> None:
    """插入占位图与图题。"""

    image_path = _ensure_placeholder(key, title, hint_lines, accent_rgb)
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)
    run = caption.add_run(f"{figure_no} {title}（占位图，建议替换为真实截图）")
    _set_run_font(run, font_name="宋体", size=10.5, color=RGBColor(74, 85, 104))


def write_chapter_1(document: Document) -> None:
    add_heading(document, "第一章 软件概述", 1)

    add_heading(document, "1.1 编写目的", 2)
    add_paragraph(document, "本文档用于说明“通用CAE分析平台软件”的安装方法、界面结构、前处理设置、求解分析流程、结果查看方式以及输入输出文件组织形式，可直接作为软件著作权申报中的《软件使用说明书》基础材料，也可作为软件交付培训文档。")

    add_heading(document, "1.2 软件简介", 2)
    add_paragraph(document, "通用CAE分析平台软件是一套面向三维有限元分析流程的桌面应用系统，采用“Python 图形界面 + C++ 求解核心 + Gmsh 网格 + PyVista/VTK 可视化”的混合架构，实现了从几何准备、材料定义、载荷边界设置、网格划分、数值求解到结果后处理和数据导出的完整工作流。")
    add_paragraph(document, "当前版本已经能够稳定支持线性静力学、非线性静力学、稳态热传导、线性模态、线性瞬态动力学和线性频响分析六类典型分析任务，适用于教学演示、科研验证、中小型结构算例测试以及桌面化 CAE 平台原型开发。")

    add_heading(document, "1.3 软件定位与适用范围", 2)
    add_bullets(
        document,
        [
            "面向对象：有限元教学用户、科研人员、算法开发者、桌面软件交付用户。",
            "问题类型：当前重点支持三维实体结构分析与稳态热传导分析。",
            "几何来源：参数化实体几何与 STEP/IGES 等 CAD 模型导入。",
            "典型规模：适合小型到中等规模模型；对大规模频响和模态问题需要结合本机内存容量评估。",
            "交付形态：支持源码运行，也支持 PyInstaller 打包后的独立可执行版本运行。",
        ],
    )

    add_heading(document, "1.4 功能概览", 2)
    add_table(
        document,
        ["模块", "当前功能", "核心实现"],
        [
            ["几何与项目管理", "参数化几何、CAD 导入、项目保存/打开、模型摘要", "PyQt6 + JSON 项目状态"],
            ["材料与工况", "材料模板、自定义材料、多材料条目管理、多载荷/多边界条目管理", "ProjectState 数据模型"],
            ["网格模块", "Gmsh 几何预览、体网格生成、质量检查、VTU 导出", "Gmsh Python API + PyVista"],
            ["求解模块", "六类分析、线性求解器切换、阻尼、线程设置", "C++/Eigen/fem_core"],
            ["结果模块", "位移/应力/应变/温度/热流显示、切片、等值面、探针、曲线、数据表", "PyVista/VTK + Matplotlib"],
            ["文档与输出", "CSV、VTU、动力曲线 CSV、频响曲线 CSV、Markdown 分析报告", "Python 导出服务"],
        ],
        widths_cm=[3.0, 8.0, 5.0],
    )

    add_heading(document, "1.5 软件技术特点", 2)
    add_bullets(
        document,
        [
            "图形界面、网格、求解和后处理分层清晰，便于后续持续扩展。",
            "求解核心采用 C++/Eigen 实现，并通过 pybind11 导出为 Python 模块 `fem_core`。",
            "几何预览与结果显示均使用三维交互视口，支持旋转、缩放、拾取和探针。",
            "分析流程中大量状态可通过 JSON 保存，便于复现实验配置。",
            "动态分析支持 Newmark-β 时间积分与 Rayleigh 阻尼，频响分析支持复数直接法。",
        ],
    )
    add_figure(
        document,
        key="fig_1_1_workflow",
        figure_no="图1-1",
        title="软件总体工作流程与数据流示意",
        hint_lines=["建议截图内容：封面页或流程图", "可展示“几何 → 网格 → 求解 → 后处理 → 导出”的完整流程"],
        accent_rgb=(30, 64, 175),
    )

    add_heading(document, "1.6 术语说明", 2)
    add_table(
        document,
        ["术语", "说明"],
        [
            ["CAD", "计算机辅助设计模型，当前可导入 STEP/STP/IGES/IGS 文件。"],
            ["Tet4", "四节点线性四面体单元，是当前结构求解的基础实体单元之一。"],
            ["Hex8", "八节点线性六面体单元，当前可在参数化长方体上直接参与求解。"],
            ["VTU", "VTK XML Unstructured Grid 格式，用于保存网格和结果场。"],
            ["Rayleigh 阻尼", "由质量矩阵和刚度矩阵线性组合得到的阻尼模型。"],
            ["Newmark-β", "线性瞬态动力学中使用的经典时间积分算法。"],
        ],
        widths_cm=[3.2, 12.0],
    )


def write_chapter_2(document: Document) -> None:
    add_heading(document, "第二章 系统要求与安装部署", 1)

    add_heading(document, "2.1 发布版运行环境要求", 2)
    add_table(
        document,
        ["项目", "建议配置"],
        [
            ["操作系统", "Windows 10/11 64 位"],
            ["处理器", "支持 64 位指令集的多核 CPU，建议 4 核及以上"],
            ["内存", "最低 8 GB，建议 16 GB 及以上"],
            ["显卡与显示", "支持 OpenGL 3.3 及以上，建议独立显卡或稳定驱动"],
            ["磁盘空间", "软件与算例建议预留 2 GB 以上空间"],
            ["运行组件", "发布版已内置 Python 运行时、Qt、VTK、Gmsh 与 fem_core 依赖"],
        ],
        widths_cm=[3.5, 11.5],
    )

    add_heading(document, "2.2 发布版目录说明", 2)
    add_paragraph(document, "当前发布版软件目录以独立文件夹形式提供，必须保持目录结构完整，不应只单独拷贝 `.exe` 文件。")
    add_table(
        document,
        ["目录或文件", "用途"],
        [
            ["CAE_Analysis_Workbench.exe", "软件主程序入口。"],
            ["_internal", "PyInstaller 打包后的运行库目录，包含 Python、Qt、VTK、Gmsh 等依赖。"],
            ["data/models", "默认网格、项目模型与中间文件输出目录。"],
            ["data/results", "默认分析结果、VTU 文件、曲线 CSV 与报告输出目录。"],
        ],
        widths_cm=[4.2, 10.8],
    )

    add_heading(document, "2.3 发布版安装与启动方式", 2)
    add_numbered(
        document,
        [
            "将完整发布目录复制到本机任意可读写位置，例如 `D:\\研究生课外\\CAE通用分析平台软件\\CAE_Analysis_Workbench`。",
            "确认文件夹保留 `CAE_Analysis_Workbench.exe`、`_internal` 与 `data` 三部分。",
            "双击 `CAE_Analysis_Workbench.exe` 启动软件，首次启动时程序会自动检查并创建 `data/models` 与 `data/results` 子目录。",
            "若软件被安全软件拦截，应将发布目录加入信任名单后重新启动。",
        ],
    )

    add_heading(document, "2.4 源码运行与开发环境说明", 2)
    add_paragraph(document, "若需要以源码方式运行或继续开发，可在本项目仓库中使用 Anaconda 环境 `pyqt6` 启动。界面层依赖 PyQt6、PyVista、VTK、Gmsh、Matplotlib、python-docx 等库；求解核心通过 `python/fem_core.cp39-win_amd64.pyd` 提供。")
    add_bullets(
        document,
        [
            "Python 解释器：`D:\\anaconda3\\envs\\pyqt6\\python.exe`。",
            "主界面入口：`python/app/launcher.py` 或 `python/app/main_window.py`。",
            "C++ 核心源码位于 `cpp/include` 与 `cpp/src`，可通过 VS2022/CMake 构建。",
            "当重新编译求解核心后，应确保新的 `fem_core` 模块位于 `python/` 目录或打包目录的运行时搜索路径内。",
        ],
    )

    add_heading(document, "2.5 初次运行说明", 2)
    add_paragraph(document, "软件首次打开时会自动初始化三维视口，并生成默认参数化长方体几何的预览。若当前显卡驱动较旧、远程桌面环境不稳定或 OpenGL 兼容性较差，首次视口初始化可能比普通窗口稍慢。")
    add_paragraph(document, "若使用打包版启动失败，可在软件目录检查是否存在 `startup_trace.log` 或 `startup_error.log`，并优先确认 `_internal` 目录未被误删、系统权限允许程序在运行目录写入 `data` 子目录。")

    add_heading(document, "2.6 数据目录与读写权限", 2)
    add_paragraph(document, "软件默认把项目、网格和分析结果写入运行目录下的 `data` 文件夹。发布版中，相对路径会自动解析到 EXE 所在目录；源码版中，相对路径会解析到仓库根目录。")
    add_bullets(
        document,
        [
            "项目文件默认保存为 `data/project_state.json` 或用户自定义 JSON 路径。",
            "网格文件默认写入 `data/models/latest_mesh.msh`。",
            "分析结果默认写入 `data/results/latest_results.csv` 与 `data/results/latest_results.vtu`。",
            "用户应避免把软件放在无写权限目录，例如某些系统保护目录或只读 U 盘根目录。",
        ],
    )


def write_chapter_3(document: Document) -> None:
    add_heading(document, "第三章 用户界面介绍", 1)

    add_heading(document, "3.1 主界面总体布局", 2)
    add_paragraph(document, "软件主界面采用典型 CAE 工作台布局，由顶部菜单栏与工具栏、左侧模型数据库、中央三维视口、右侧属性编辑器以及底部消息日志组成。用户可以通过左侧模块导航切换到“几何、材料、载荷与边界、热模块、网格、分析步、结果、视图”等页面。")
    add_figure(
        document,
        key="fig_3_1_main_ui",
        figure_no="图3-1",
        title="主界面总体布局",
        hint_lines=["建议截图内容：完整主界面", "请同时包含左侧树、中央视口、右侧属性编辑器和底部日志"],
        accent_rgb=(37, 99, 235),
    )

    add_heading(document, "3.2 菜单栏功能", 2)
    add_bullets(
        document,
        [
            "文件：导入 CAD、打开项目、保存项目、查看模型摘要。",
            "视图：等轴测、前视、右视、俯视、适配视图。",
            "网格：生成网格、检查质量、导出网格。",
            "求解：运行分析、结果探针、结果数据表、导出报告、导出结果 CSV、导出结果 VTU。",
        ],
    )

    add_heading(document, "3.3 主工具栏说明", 2)
    add_paragraph(document, "主工具栏提供最常用工作流的快捷按钮，包括导入 CAD、预览几何、生成网格、检查质量、运行分析、结果探针、结果表格、导出报告和模型摘要等操作。该设计可以缩短常规流程下的操作路径。")

    add_heading(document, "3.4 视图工具栏说明", 2)
    add_paragraph(document, "视图工具栏用于快速控制三维相机姿态和辅助显示，包括等轴视图、前视图、右视图、俯视图、适配视图以及坐标轴显示开关。对于几何选面、网格检查和结果后处理观察尤为常用。")

    add_heading(document, "3.5 左侧模型数据库", 2)
    add_paragraph(document, "左侧树形区域以模块方式组织当前项目。用户单击节点即可切换右侧属性编辑器内容，不同页面的参数会同步写入统一的项目状态对象。该区域本质上相当于轻量级模型数据库，适合记录当前项目各模块配置。")

    add_heading(document, "3.6 右侧属性编辑器", 2)
    add_paragraph(document, "右侧属性编辑器是软件的核心操作区。不同模块会动态显示对应参数表单，例如几何尺寸、材料属性、载荷边界、网格算法、求解器参数和结果后处理设置等。")
    add_paragraph(document, "需要特别注意：材料、载荷、边界、热边界和热载荷等页面均采用“修改参数后点击应用”的确认机制，防止用户在未确认配置的情况下误用旧工况求解。")

    add_heading(document, "3.7 中央三维视口", 2)
    add_paragraph(document, "中央视口基于 PyVista/VTK 构建，可显示几何预览、体网格、结果云图、切片、等值面和探针高亮等内容。不同场景下，视口会自动切换标题与提示信息，以便用户识别当前正在查看的是几何、网格还是结果。")

    add_heading(document, "3.8 底部消息日志与状态栏", 2)
    add_paragraph(document, "底部消息日志会实时记录后台任务执行情况，例如几何预览、网格生成、质量检查、求解运行、导出文件和图形选面提示等；状态栏则用于显示当前忙碌状态和短时消息。")
    add_figure(
        document,
        key="fig_3_2_toolbar_detail",
        figure_no="图3-2",
        title="菜单栏、工具栏与日志区域细节",
        hint_lines=["建议截图内容：顶部菜单栏和工具栏", "可配合底部日志输出一起展示操作反馈效果"],
        accent_rgb=(14, 116, 144),
    )


def write_chapter_4(document: Document) -> None:
    add_heading(document, "第四章 项目管理", 1)

    add_heading(document, "4.1 新建项目与默认状态", 2)
    add_paragraph(document, "软件启动后会自动创建一个默认项目，默认名称为“新建CAE分析项目”，并初始化参数化长方体几何、结构钢材料、默认载荷与边界条件、线性静力学分析步以及默认结果输出路径。")

    add_heading(document, "4.2 导入外部 CAD 模型", 2)
    add_paragraph(document, "通过“文件 → 导入 CAD”或主工具栏按钮可导入 `*.step`、`*.stp`、`*.iges`、`*.igs` 文件。导入后，几何来源会切换为“外部 CAD”，中央视口刷新为 CAD 轮廓的几何预览，原有网格和结果会被清空。")
    add_paragraph(document, "导入 CAD 后，参数化尺寸输入框不再代表实际外部几何，只用于参数化模式下的备用状态。因此在 CAD 模式下，应以预览图与后续生成网格结果为准。")
    add_figure(
        document,
        key="fig_4_1_project_save_load",
        figure_no="图4-1",
        title="项目保存、打开与模型摘要",
        hint_lines=["建议截图内容：保存项目、打开项目、模型摘要三个功能入口", "也可展示保存后的 JSON 文件示例位置"],
        accent_rgb=(22, 163, 74),
    )

    add_heading(document, "4.3 恢复为参数化几何", 2)
    add_paragraph(document, "若用户希望从外部 CAD 模式返回参数化建模，可在几何页面点击“恢复参数化”。该操作会把几何模式重置为默认参数化长方体，同时清除当前已有网格与结果。")

    add_heading(document, "4.4 保存项目", 2)
    add_paragraph(document, "项目通过 JSON 文件保存。执行“保存项目”后，软件会把当前 `ProjectState` 中的几何、材料、工况、网格摘要、求解参数和结果摘要等信息序列化输出。")
    add_paragraph(document, "需要注意的是，项目文件保存的是状态和摘要，而不是完整体网格与完整结果场。因此项目重开后通常需要重新生成体网格，并在必要时重新运行分析以恢复三维结果场。")

    add_heading(document, "4.5 打开项目", 2)
    add_paragraph(document, "执行“打开项目”后，软件会根据 JSON 内容重建项目状态，并同步刷新界面控件、几何预览和模型摘要。出于一致性考虑，加载项目后当前网格对象、质量检查结果和分析对象会被重置为空。")

    add_heading(document, "4.6 模型摘要窗口", 2)
    add_paragraph(document, "模型摘要窗口用于集中浏览当前项目的关键参数，包括几何类型、网格算法、材料信息、荷载边界、求解配置以及结果摘要。该窗口适合用于求解前的配置核查，也适合汇总写入报告。")
    add_figure(
        document,
        key="fig_4_2_cad_switch",
        figure_no="图4-2",
        title="CAD 导入与参数化几何切换",
        hint_lines=["建议截图内容：几何页面中的“导入 CAD”“恢复参数化”按钮", "可在图中展示当前几何来源状态标签"],
        accent_rgb=(132, 90, 223),
    )


def write_chapter_5(document: Document) -> None:
    add_heading(document, "第五章 前处理：几何与材料设置", 1)

    add_heading(document, "5.1 参数化几何类型", 2)
    add_paragraph(document, "当前参数化几何模块支持四类基础实体：长方体、圆柱体、球体和带孔板。不同几何类型会动态启用对应尺寸参数，例如长方体使用长度/宽度/高度，圆柱体使用长度和半径，带孔板使用长度、宽度、高度和孔半径。")

    add_heading(document, "5.2 几何预览", 2)
    add_paragraph(document, "几何预览使用 Gmsh 的二维表面网格快速生成外形轮廓，目的是帮助用户在生成三维体网格前检查几何来源、尺寸和可选表面。该预览不会直接参与求解，但会为后续图形选面提供真实几何面补丁。")
    add_figure(
        document,
        key="fig_5_1_geometry_page",
        figure_no="图5-1",
        title="参数化几何设置页",
        hint_lines=["建议截图内容：几何页完整界面", "包含参数化类型、尺寸输入、CAD 导入与几何预览按钮"],
        accent_rgb=(37, 99, 235),
    )

    add_heading(document, "5.3 材料模板与自定义材料", 2)
    add_paragraph(document, "材料页面内置三种模板：结构钢、铝合金和混凝土。选择模板后，软件会自动填充杨氏模量、泊松比、密度、屈服强度和硬化模量等推荐值，用户可进一步修改并保存为自定义材料名称。")
    add_paragraph(document, "材料页面支持新增材料条目和删除材料条目，因此一个项目中可以维护材料库；但当前求解时仍以“当前激活材料”赋予整个几何体。该点属于现阶段实现策略，已在第十二章中列为限制说明。")

    add_heading(document, "5.4 材料参数说明", 2)
    add_bullets(
        document,
        [
            "杨氏模量 E [Pa]：控制线弹性刚度。",
            "泊松比 ν：控制横向变形耦合关系，当前输入范围为 0.01 到 0.49。",
            "密度 ρ [kg/m³]：参与模态、瞬态和频响分析中的质量矩阵组装。",
            "屈服强度 [Pa]：当启用材料非线性时，用于判断是否进入屈服。",
            "硬化模量 [Pa]：当材料进入屈服后，用于形成简化的等向强化切线模量。",
        ],
    )

    add_heading(document, "5.5 材料库管理", 2)
    add_paragraph(document, "新增材料时，软件会以当前激活材料为模板复制参数，再要求用户输入新名称；删除材料时，软件要求至少保留一个材料条目，以避免项目进入非法状态。材料参数修改后不会立即生效，必须点击“应用材料设置”。")

    add_heading(document, "5.6 热材料参数与当前适用方式", 2)
    add_paragraph(document, "热模块与材料模块共享导热系数 `k [W/(m·K)]`。稳态热传导分析中，会使用当前激活材料的导热系数组装热传导矩阵。当前版本中，热材料仍采用“一个激活材料作用于整个当前几何体”的方式。")
    add_figure(
        document,
        key="fig_5_2_material_page",
        figure_no="图5-2",
        title="材料参数与热材料设置页",
        hint_lines=["建议截图内容：材料页面", "请包含材料模板、自定义材料、非线性复选框和热导率输入"],
        accent_rgb=(22, 163, 74),
    )


def write_chapter_6(document: Document) -> None:
    add_heading(document, "第六章 载荷、边界与热条件设置", 1)

    add_heading(document, "6.1 荷载工况设置", 2)
    add_paragraph(document, "载荷页面允许用户维护多个荷载工况，并通过“当前工况”下拉框切换激活项。每个荷载工况包含 `Fx`、`Fy`、`Fz` 三个力分量、受载面以及识别容差。求解时总载荷会被平均分配到选定受载面的所有节点。")

    add_heading(document, "6.2 图形选择受载面", 2)
    add_paragraph(document, "点击“图形选择受载面”后，三维视口会进入选面模式。软件优先基于真实几何面补丁进行拾取，并以高亮颜色标识鼠标悬停面和已选择面。用户采用左键单击释放完成选择，拖动旋转视图不会触发误选。")

    add_heading(document, "6.3 边界条件设置", 2)
    add_paragraph(document, "边界条件页面支持多个边界条目，每个条目可分别设置 X/Y/Z 三个方向是否约束，以及对应位移值。当前界面既支持固定约束（位移值为 0），也支持给定位移边界。")

    add_heading(document, "6.4 图形选择约束面", 2)
    add_paragraph(document, "约束面选择与受载面选择采用同一套图形拾取机制。若真实表面补丁可用，则优先使用 `surface:编号` 节点集；若仅使用兼容端面名称，则求解时会基于几何包围盒极值和容差自动识别 `xmin/xmax/ymin/ymax/zmin/zmax` 节点。")

    add_heading(document, "6.5 固定温度边界", 2)
    add_paragraph(document, "热模块中的“固定温度边界”用于指定一个表面上的定温约束。界面提供图形选面按钮和温度输入框。当前该设置表示稳态热传导问题中的边界温度，并不表示整个模型的初始温度场。")

    add_heading(document, "6.6 热流载荷", 2)
    add_paragraph(document, "热流载荷通过总热流功率 `W` 输入，并在求解时平均分配到所选表面的所有节点。若热流值为 0，允许不选择热流面；若热流值非 0，则必须先选择热流作用面再点击应用。")
    add_figure(
        document,
        key="fig_6_1_load_boundary",
        figure_no="图6-1",
        title="力载荷与位移边界设置页",
        hint_lines=["建议截图内容：载荷与边界页面", "包含当前工况、多边界条目、力分量输入和图形选面按钮"],
        accent_rgb=(220, 38, 38),
    )

    add_heading(document, "6.7 应用确认机制与容差设置", 2)
    add_paragraph(document, "载荷、边界、温度边界和热流载荷均需要点击对应“应用”按钮后才视为已确认，软件内部会将 `is_applied` 标记置为 `True`。一旦参数再次修改，标记会自动重置，提醒用户在运行分析前重新确认。")
    add_paragraph(document, "“边界识别容差比例”用于控制兼容端面识别时的节点筛选范围。模型尺寸跨度越大或 CAD 表面越不规则，越需要谨慎设置容差，以避免选面过窄或误选过宽。")
    add_figure(
        document,
        key="fig_6_2_thermal_page",
        figure_no="图6-2",
        title="固定温度边界与热流载荷设置页",
        hint_lines=["建议截图内容：热模块页面", "请包含图形选择温度边界、图形选择热流面和温度/热流输入项"],
        accent_rgb=(249, 115, 22),
    )


def write_chapter_7(document: Document) -> None:
    add_heading(document, "第七章 网格划分与质量检查", 1)

    add_heading(document, "7.1 网格控制参数", 2)
    add_table(
        document,
        ["参数", "作用", "当前说明"],
        [
            ["网格拓扑", "选择四面体或六面体", "六面体当前仅支持参数化长方体"],
            ["全局网格尺寸", "控制整体单元尺寸", "数值越小，单元越密，计算规模越大"],
            ["2D 网格算法", "控制几何预览与面网格生成", "支持 Delaunay、Frontal-Delaunay、BAMG"],
            ["3D 网格算法", "控制体网格生成", "支持 Delaunay、Frontal、MMG3D、HXT"],
            ["单元阶次", "Gmsh 网格阶次", "界面允许 1 或 2 阶，当前求解以内角点线性单元为主"],
            ["Gmsh 网格优化", "生成后调用优化流程", "建议默认开启"],
            ["局部加密", "在受载端附近设置盒状背景尺寸场", "当前仅对四面体网格有效"],
        ],
        widths_cm=[3.0, 4.0, 8.0],
    )

    add_heading(document, "7.2 支持的网格类型与单元阶次", 2)
    add_paragraph(document, "当前软件在界面上支持四面体网格和六面体网格两种拓扑。对于参数化长方体，六面体网格使用结构化 transfinite/recombine 控制直接生成 Hex 网格，并可由求解器按 Hex8 单元参与计算。")
    add_paragraph(document, "界面还支持 1 阶和 2 阶 Gmsh 元素阶次，但当前求解内核在结构与热传导计算中主要读取角点节点，因此二阶单元会退化为角点线性单元参与求解。这一点请在正式工程使用前充分评估。")

    add_heading(document, "7.3 生成体网格", 2)
    add_paragraph(document, "点击“生成体网格”后，软件会基于当前几何、网格算法、局部加密设置和拓扑选项调用 Gmsh 生成三维体网格，并自动输出 `latest_mesh.msh`。生成成功后，中央视口切换为体网格显示，模型摘要同步写入节点数、显示单元数、求解单元数和单元类型。")

    add_heading(document, "7.4 网格质量检查", 2)
    add_paragraph(document, "网格质量检查使用 PyVista 的 `scaled_jacobian` 指标，并给出最小值、平均值、最大值以及低质量单元统计。其中，程序把 `<=0.2` 视为较差单元，把 `(0.2, 0.4]` 视为预警单元。")
    add_paragraph(document, "检查结果既会显示在模型摘要中，也会用于帮助用户判断是否需要调小网格尺寸、切换算法或清理几何。")
    add_figure(
        document,
        key="fig_7_1_mesh_control",
        figure_no="图7-1",
        title="网格控制参数设置页",
        hint_lines=["建议截图内容：网格页面", "请包含拓扑、网格尺寸、算法、单元阶次和局部加密参数"],
        accent_rgb=(79, 70, 229),
    )

    add_heading(document, "7.5 网格导出", 2)
    add_paragraph(document, "当前网格可导出为 VTU 文件，用于与 ParaView 等第三方可视化软件联动。导出的对象为界面显示网格；若当前生成的是六面体网格，导出的也是原始 Hex 显示网格。")

    add_heading(document, "7.6 使用建议", 2)
    add_bullets(
        document,
        [
            "几何预览正常后再生成体网格，可减少 CAD 几何错误导致的失败。",
            "对规则块体优先尝试六面体网格，对复杂外部 CAD 优先使用四面体网格。",
            "启用局部加密时建议先选择受载面，否则加密区域可能不生效。",
            "若出现质量较差单元，可尝试减小网格尺寸、切换 3D 算法或简化 CAD 几何。",
            "若计划进行模态、频响等矩阵更重的分析，应避免过密网格。",
        ],
    )
    add_figure(
        document,
        key="fig_7_2_mesh_quality",
        figure_no="图7-2",
        title="网格质量检查结果与体网格显示",
        hint_lines=["建议截图内容：体网格显示效果和质量检查摘要", "可在图中同时体现消息日志中的质量统计"],
        accent_rgb=(14, 165, 233),
    )


def write_chapter_8(document: Document) -> None:
    add_heading(document, "第八章 求解分析功能", 1)

    add_heading(document, "8.1 分析类型总览", 2)
    add_table(
        document,
        ["分析类型", "物理问题", "核心算法", "主要结果"],
        [
            ["线性静力学", "小变形静力平衡", "组装 K、施加载荷与罚函数边界后求解 KU=F", "位移、Von Mises 应力、等效应变"],
            ["非线性静力学", "材料非线性增量静力", "载荷步 + Newton-Raphson + 切线刚度迭代", "位移、应力、应变、迭代次数、残差"],
            ["稳态热传导", "定常导热", "组装热传导矩阵 Kt，求解 KtT=Q", "温度、热流密度"],
            ["线性模态分析", "固有频率与振型", "自由度缩减后的广义特征值求解 Kφ=λMφ", "模态频率、振型"],
            ["线性瞬态动力学", "时域动力响应", "Newmark-β 时间积分 + Rayleigh 阻尼", "最终时刻结果、全局最大位移时程、受载面平均响应"],
            ["线性频响分析", "稳态谐响应", "逐频率求解复数动态刚度 [K+iωC-ω²M]u=F", "峰值响应场、幅频曲线、峰值频率"],
        ],
        widths_cm=[3.0, 3.5, 6.0, 4.0],
    )

    add_heading(document, "8.2 线性求解器与并行配置", 2)
    add_paragraph(document, "对线性静力学、非线性静力学、线性瞬态动力学和稳态热传导四类分析，用户可以在界面中选择线性方程组求解器。模态分析和频响分析使用各自内置算法，不调用该线性求解器选项。")
    add_table(
        document,
        ["求解器", "适用分析", "特点与建议"],
        [
            ["SparseLU 直接法", "线性静力、非线性迭代、瞬态、热传导", "数值稳定，适合中小规模模型，内存消耗相对较高。"],
            ["Conjugate Gradient", "线性静力、非线性迭代、瞬态、热传导", "迭代法，适合条件较好的对称正定问题，内存占用较低。"],
            ["BiCGSTAB", "线性静力、非线性迭代、瞬态、热传导", "对更一般的稀疏线性系统较稳健，可作为迭代求解备选。"],
        ],
        widths_cm=[4.0, 4.5, 8.0],
    )
    add_paragraph(document, "并行设置当前仅作用于 CPU OpenMP 路径，主要体现在全局矩阵组装与部分后处理阶段。当前版本尚未接入 CUDA GPU 求解或 MPI 分布式求解。")

    add_heading(document, "8.3 线性静力学分析", 2)
    add_paragraph(document, "线性静力学分析是软件的默认分析类型，适用于小变形、线弹性结构问题。计算流程为：根据当前网格与材料组装全局刚度矩阵 `K`，根据受载面将总力平均分配到节点形成 `F`，再通过罚函数法施加位移边界，最终求解 `KU=F`。")
    add_paragraph(document, "求解完成后，软件会提取节点位移、单元 Von Mises 等效应力和等效应变，并生成结果 CSV 与 VTU 文件。")

    add_heading(document, "8.4 非线性静力学分析", 2)
    add_paragraph(document, "非线性静力学分析采用“载荷步 + Newton-Raphson 迭代”流程。用户可设置载荷步数、每步最大迭代次数和收敛容差。每个载荷步中，程序会根据当前位移场重新组装切线刚度矩阵，并用残差范数判断是否收敛。")
    add_paragraph(document, "当前材料非线性实现为简化的等向强化模型：当等效应变超过 `屈服强度 / E` 后，材料使用由 `E` 与硬化模量计算得到的等效塑性切线模量，并对等效应力进行比例缩放。该实现适合教学和算法验证，不应直接替代成熟工业弹塑性本构。")

    add_heading(document, "8.5 稳态热传导分析", 2)
    add_paragraph(document, "稳态热传导分析使用当前激活材料的导热系数 `k` 组装热传导矩阵，对固定温度边界采用罚函数法施加，对热流面采用节点平均热功率加载，最终求解稳态温度场。")
    add_paragraph(document, "结果包括节点温度、单元热流密度幅值以及导出的温度/热流 CSV 和 VTU 文件。热分析结果页面只允许显示温度和热流密度两种结果量。")

    add_heading(document, "8.6 线性模态分析", 2)
    add_paragraph(document, "线性模态分析会先组装结构刚度矩阵和质量矩阵，再去除被边界条件约束的自由度，构造自由自由度子系统，并调用 `Eigen::GeneralizedSelfAdjointEigenSolver` 求解广义特征值问题。")
    add_paragraph(document, "输出结果包括多阶固有频率及对应振型。当前界面默认显示第一阶振型，用户可在结果页面切换到指定阶次。由于当前实现使用自由度缩减后的稠密特征值求解，模型规模过大时会明显增加内存压力。")

    add_heading(document, "8.7 线性瞬态动力学分析", 2)
    add_paragraph(document, "线性瞬态动力学分析使用 Newmark-β 时间积分算法。程序会组装刚度矩阵 `K`、质量矩阵 `M` 和阻尼矩阵 `C=αM+βK`，然后按照用户设置的总时长、时间步长、β 和 γ 系数逐步推进时域响应。")
    add_paragraph(document, "软件会记录整个时程中的全局最大位移、受载面平均 `Ux/Uy/Uz/|U|` 响应，并在结果页面提供摘要窗口和曲线导出功能。当前质量矩阵采用集中质量形式，更适合教学演示与中小型模型快速响应分析。")

    add_heading(document, "8.8 线性频响分析", 2)
    add_paragraph(document, "线性频响分析按频率点逐点计算稳态谐响应。程序在自由度缩减后构造复数动态刚度矩阵 `[K + iωC - ω²M]`，使用复数直接求解得到每个频率点的响应幅值，并记录全局最大位移幅值与受载面平均响应幅值。")
    add_paragraph(document, "界面中显示的三维结果场对应于全频率扫描过程中峰值响应最大的那一个频率点。对于频响曲线本身，用户可切换查看全局最大位移、受载面平均 `Ux/Uy/Uz/|U|` 幅值。")
    add_paragraph(document, "由于当前实现对自由度子系统使用稠密复数矩阵逐频率求解，因此更适合小型到中等规模模型的频域研究。")

    add_heading(document, "8.9 阻尼设置", 2)
    add_bullets(
        document,
        [
            "无阻尼：动态分析中不加入 Rayleigh 阻尼项。",
            "直接 Rayleigh 阻尼：用户直接输入 `α` 和 `β`，程序按 `C=αM+βK` 组装阻尼矩阵。",
            "按阻尼比换算 Rayleigh：用户给定目标阻尼比 `ζ` 以及两个参考频率 `f1、f2`，程序自动换算出 `α` 和 `β`。",
        ],
    )
    add_paragraph(document, "当选择“按阻尼比换算 Rayleigh”时，两个参考频率必须均大于 0，且不能相同，否则程序会直接报错。")

    add_heading(document, "8.10 求解流程、日志与失败处理", 2)
    add_paragraph(document, "点击“运行分析”后，软件会在后台线程中执行求解任务，避免主界面阻塞。运行期间，主工具栏、属性编辑器和菜单会被临时禁用，状态栏显示忙碌信息。")
    add_paragraph(document, "若求解失败，底部消息日志会输出完整异常栈，典型原因包括：未应用材料设置、未应用载荷或边界、未识别到目标节点、非线性迭代未收敛、网格退化或热边界未正确选择。")
    add_figure(
        document,
        key="fig_8_1_solver_setup",
        figure_no="图8-1",
        title="分析步、求解器、阻尼与线程设置页",
        hint_lines=["建议截图内容：分析步页面", "需覆盖分析类型、线性求解器、模态阶数、Newmark 参数、频率范围、阻尼和线程数"],
        accent_rgb=(37, 99, 235),
    )
    add_figure(
        document,
        key="fig_8_2_solver_log",
        figure_no="图8-2",
        title="求解运行过程与消息日志反馈",
        hint_lines=["建议截图内容：分析完成后的日志输出", "可展示静力、热分析或动态分析完成后的关键摘要消息"],
        accent_rgb=(22, 163, 74),
    )


def write_chapter_9(document: Document) -> None:
    add_heading(document, "第九章 后处理与结果查看", 1)

    add_heading(document, "9.1 结果显示模式", 2)
    add_paragraph(document, "结果页面会根据当前分析类型自动刷新可选结果量。结构分析默认支持位移、应力和应变；稳态热传导支持温度和热流密度；模态结果仅支持位移振型显示。")

    add_heading(document, "9.2 整体云图与变形显示", 2)
    add_paragraph(document, "“显示整体结果”会在中央视口中加载完整结果云图，并按当前结果量应用颜色映射。结构问题默认叠加变形显示，变形倍数可由用户输入调整；热分析结果则不施加几何变形。")
    add_figure(
        document,
        key="fig_9_1_result_overall",
        figure_no="图9-1",
        title="整体结果云图显示",
        hint_lines=["建议截图内容：完整结果云图", "可选位移云图、Von Mises 应力云图或温度云图"],
        accent_rgb=(220, 38, 38),
    )

    add_heading(document, "9.3 结果切片", 2)
    add_paragraph(document, "切片功能支持按 X、Y、Z 三个方向生成单一平面切片，并允许用户通过百分比控制切片位置。对结构结果可选择在变形后结果上切片，也可在原始结果网格上切片。")

    add_heading(document, "9.4 结果等值面", 2)
    add_paragraph(document, "等值面功能会对当前结果量生成多层等值面，层数范围为 2 到 200。若结果量几乎没有变化，或当前分析类型不支持相应标量，程序会给出明确提示。")
    add_figure(
        document,
        key="fig_9_2_result_slice_contour",
        figure_no="图9-2",
        title="结果切片与等值面显示",
        hint_lines=["建议截图内容：切片和等值面两种效果", "建议在同一页展示切片方向、位置和等值面层数设置"],
        accent_rgb=(249, 115, 22),
    )

    add_heading(document, "9.5 模态振型查看", 2)
    add_paragraph(document, "当分析类型为模态分析时，结果页面会显示模态阶次输入框和“显示选定模态”按钮。用户选择某一阶后，软件会切换到对应振型，并在日志中输出该阶固有频率。")

    add_heading(document, "9.6 动力响应摘要与曲线", 2)
    add_paragraph(document, "瞬态动力学分析完成后，软件会提供“查看动力响应摘要”和“导出动力曲线 CSV”功能。曲线可选择显示全局最大位移、受载面平均 `Ux/Uy/Uz/|U|` 五种响应量。")

    add_heading(document, "9.7 频响曲线查看", 2)
    add_paragraph(document, "频响分析完成后，软件会启用“查看频响曲线”和“导出频响曲线 CSV”功能。曲线横坐标为频率，纵坐标为用户选定响应量的幅值。")
    add_figure(
        document,
        key="fig_9_3_result_probe_table",
        figure_no="图9-3",
        title="结果探针与结果数据表",
        hint_lines=["建议截图内容：探针窗口和结果表格窗口", "可体现节点坐标、位移、应力、温度等查询信息"],
        accent_rgb=(22, 163, 74),
    )

    add_heading(document, "9.8 结果探针", 2)
    add_paragraph(document, "结果探针开启后，用户可在结果表面左键单击任意位置，程序会先用 CellPicker 拾取世界坐标，再吸附到最近节点，并弹出探针窗口显示节点坐标、位移分量、位移幅值、应力、应变、温度和热流密度等信息。")
    add_paragraph(document, "为减少误操作，探针与选面模块共用“单击释放判定”机制，拖动画面不会触发探针查询。")

    add_heading(document, "9.9 结果数据表", 2)
    add_paragraph(document, "结果数据表窗口可以浏览多类结构化表格，包括：结果摘要、节点结果预览、单元结果预览、模态频率表、瞬态响应时程表、频响响应表。每个表都可以单独导出为 CSV。")

    add_heading(document, "9.10 结果文件导出与分析报告", 2)
    add_bullets(
        document,
        [
            "导出结果 CSV：结构分析由求解核心直接导出节点位移与单元结果；热分析导出温度与热流数据。",
            "导出结果 VTU：导出当前完整结果网格，便于第三方可视化软件继续处理。",
            "导出动力曲线 CSV：导出瞬态分析的时间历程曲线。",
            "导出频响曲线 CSV：导出频响分析的幅频曲线。",
            "导出分析报告：当前导出 Markdown 文档，内容包括几何、网格、材料、工况、求解配置和结果摘要。",
        ],
    )
    add_figure(
        document,
        key="fig_9_4_dynamic_curves",
        figure_no="图9-4",
        title="动力响应摘要与频响曲线窗口",
        hint_lines=["建议截图内容：瞬态响应曲线窗口或频响曲线窗口", "可在图中体现曲线类型切换与导出按钮"],
        accent_rgb=(79, 70, 229),
    )


def write_chapter_10(document: Document) -> None:
    add_heading(document, "第十章 典型使用流程", 1)

    add_heading(document, "10.1 线性静力学典型流程", 2)
    add_numbered(
        document,
        [
            "启动软件并检查默认几何预览是否正常。",
            "在几何页面选择参数化类型或导入外部 CAD。",
            "在材料页面选择材料模板并点击“应用材料设置”。",
            "在载荷页面设置力分量，图形选择受载面并应用载荷。",
            "在边界页面图形选择约束面，设置约束方向并应用边界。",
            "在网格页面设置网格尺寸与算法，生成体网格并检查质量。",
            "在分析步页面选择“线性静力学”和合适的线性求解器。",
            "点击“运行分析”，完成后在结果页查看位移、应力和应变。",
            "必要时导出 CSV、VTU 和分析报告。",
        ],
    )

    add_heading(document, "10.2 稳态热传导典型流程", 2)
    add_numbered(
        document,
        [
            "准备几何并设置材料导热系数。",
            "生成体网格，确认网格质量满足要求。",
            "在热模块中图形选择固定温度边界，输入温度并点击“应用温度边界”。",
            "若存在热输入，再图形选择热流面，输入总热流功率并点击“应用热流载荷”。",
            "在分析步页面选择“稳态热传导”，设置热方程求解器。",
            "运行分析并在结果页查看温度云图、热流密度、结果表和 VTU 文件。",
        ],
    )

    add_heading(document, "10.3 动力分析典型流程", 2)
    add_numbered(
        document,
        [
            "完成几何、材料、载荷、边界和网格设置。",
            "若进行模态分析，建议先检查边界是否足以消除刚体模态。",
            "若进行瞬态分析，设置总时长、时间步长、Newmark β/γ 和阻尼模式。",
            "若进行频响分析，设置起始频率、终止频率、频率点数和阻尼模式。",
            "运行分析后，在结果页查看模态振型、动力响应摘要或频响曲线。",
            "根据需要导出时程 CSV、幅频 CSV、结果 VTU 和 Markdown 报告。",
        ],
    )
    add_figure(
        document,
        key="fig_10_1_typical_flow",
        figure_no="图10-1",
        title="典型分析操作流程示意",
        hint_lines=["建议截图内容：静力或热分析的全流程组合图", "也可用一张整合排版图展示从建模到后处理的步骤"],
        accent_rgb=(30, 64, 175),
    )


def write_chapter_11(document: Document) -> None:
    add_heading(document, "第十一章 输入输出文件与数据说明", 1)

    add_heading(document, "11.1 输入文件", 2)
    add_table(
        document,
        ["文件类型", "扩展名", "说明"],
        [
            ["CAD 模型", "*.step / *.stp / *.iges / *.igs", "用于外部几何导入。"],
            ["项目文件", "*.json", "保存项目状态、求解参数和摘要信息。"],
            ["源码运行入口", "*.py", "用于开发模式下启动界面或辅助生成文档。"],
        ],
        widths_cm=[3.2, 5.2, 8.0],
    )

    add_heading(document, "11.2 中间文件与运行时文件", 2)
    add_table(
        document,
        ["文件或目录", "默认位置", "作用"],
        [
            ["网格文件", "data/models/latest_mesh.msh", "Gmsh 生成后的体网格文件。"],
            ["启动轨迹", "startup_trace.log", "仅在诊断模式下用于记录打包版启动轨迹。"],
            ["消息日志", "界面底部实时显示", "记录当前运行状态，不单独生成文件。"],
        ],
        widths_cm=[3.0, 5.8, 7.6],
    )

    add_heading(document, "11.3 结果输出文件", 2)
    add_table(
        document,
        ["文件", "默认位置", "内容"],
        [
            ["结构分析 CSV", "data/results/latest_results.csv", "节点位移、单元应力和应变等结果。"],
            ["热分析 CSV", "data/results/latest_results.csv", "节点温度与单元热流密度。"],
            ["结果 VTU", "data/results/latest_results.vtu", "完整结果场网格。"],
            ["瞬态曲线 CSV", "data/results/transient_history.csv", "时间-响应值曲线。"],
            ["频响曲线 CSV", "data/results/frequency_response.csv", "频率-幅值曲线。"],
            ["分析报告", "用户指定的 *.md", "Markdown 格式结果报告。"],
        ],
        widths_cm=[3.2, 5.8, 7.4],
    )

    add_heading(document, "11.4 项目文件说明", 2)
    add_paragraph(document, "项目 JSON 文件包含几何参数、材料库、荷载工况库、边界条件库、热工况、求解器设置、网格摘要和结果摘要等字段。它能够恢复“配置状态”，但不会自动恢复完整的网格对象和三维结果场对象。")


def write_chapter_12(document: Document) -> None:
    add_heading(document, "第十二章 注意事项、常见问题与已知限制", 1)

    add_heading(document, "12.1 使用注意事项", 2)
    add_bullets(
        document,
        [
            "修改材料、载荷、边界或热条件后，请重新点击对应“应用”按钮。",
            "运行分析前请先生成体网格；仅有几何预览时无法进入求解。",
            "若使用 CAD 导入模型，建议先检查几何预览是否闭合、是否存在异常退化面。",
            "若频响或模态分析计算规模较大，建议适当减小网格密度或降低模态阶数/频率点数。",
            "发布版软件目录中的 `_internal` 属于运行必需目录，不得删除。",
        ],
    )

    add_heading(document, "12.2 常见问题处理", 2)
    add_table(
        document,
        ["现象", "可能原因", "处理建议"],
        [
            ["提示“请先生成网格”", "尚未生成三维体网格", "先进入网格模块执行“生成体网格”。"],
            ["提示“材料未确认”", "修改后未点击“应用材料设置”", "返回材料页面完成应用。"],
            ["提示“未识别到受力端节点/固定端节点”", "选面未正确设置或容差不合适", "重新图形选面，必要时调整边界识别容差比例。"],
            ["非线性分析未收敛", "载荷过大、步数过少、容差过严或网格质量不佳", "增加载荷步数，优化网格，检查材料参数。"],
            ["热分析无热流结果", "热流功率为 0 或未应用热流载荷", "检查热流输入值和热流面设置。"],
            ["打包版无法启动", "目录缺失、权限不足或图形环境异常", "检查 `_internal` 与 `data` 是否完整，并查看启动日志。"],
        ],
        widths_cm=[4.0, 5.2, 7.2],
    )

    add_heading(document, "12.3 已知限制", 2)
    add_bullets(
        document,
        [
            "当前多材料库已支持界面管理，但求解时仍将“当前激活材料”赋予整个当前几何体。",
            "界面允许生成二阶网格，但当前求解核心主要使用角点节点，即按线性 Tet4/Hex8 单元求解。",
            "六面体网格当前仅稳定支持参数化长方体，且不支持局部加密。",
            "模态分析和频响分析基于自由度缩减后的稠密矩阵算法，更适合小型到中等规模模型。",
            "当前动态分析为线性动力学流程，不包含接触、几何大变形和非线性瞬态。",
            "结果报告当前导出为 Markdown，不直接生成 PDF/Word 格式。",
            "并行能力当前限于 CPU OpenMP 路径，未集成 GPU 或分布式并行。",
        ],
    )

    add_heading(document, "12.4 软著整理与截图替换建议", 2)
    add_paragraph(document, "在用于软件著作权申报时，建议将本文档中的占位图逐一替换为真实运行截图，并保持图题与正文说明一致。优先补充以下截图：主界面总览、几何设置、材料设置、载荷边界、热模块、网格设置、求解设置、结果云图、探针查询和典型流程。")
    add_paragraph(document, "替换截图时建议保留当前图号和图题不变，只更新图片内容，这样可以避免目录、交叉引用和页码重新大范围调整。")


def build_manual(output_path: Path, title: str, version: str) -> None:
    """构建并保存 Word 文档。"""

    document = Document()
    _style_document(document)
    add_cover(document, title=title, version=version)
    add_toc(document)

    write_chapter_1(document)
    write_chapter_2(document)
    write_chapter_3(document)
    write_chapter_4(document)
    write_chapter_5(document)
    write_chapter_6(document)
    write_chapter_7(document)
    write_chapter_8(document)
    write_chapter_9(document)
    write_chapter_10(document)
    write_chapter_11(document)
    write_chapter_12(document)

    for section in document.sections:
        add_footer_page_number(section)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def backup_existing_file(output_path: Path) -> Path | None:
    """在覆盖前备份旧文档。"""

    if not output_path.exists():
        return None
    backup_name = f"{output_path.stem}_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}{output_path.suffix}"
    backup_path = output_path.with_name(backup_name)
    shutil.copy2(output_path, backup_path)
    return backup_path


def inspect_manual(output_path: Path) -> dict[str, int]:
    """对生成后的文档做基本自检。"""

    document = Document(output_path)
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    table_text: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text)
    full_text = "\n".join(paragraph_text + table_text)
    required_keywords = [
        "SparseLU",
        "Conjugate Gradient",
        "BiCGSTAB",
        "Newton-Raphson",
        "Newmark",
        "Rayleigh",
        "频响",
        "模态",
        "Tet4",
        "Hex8",
    ]
    missing = [keyword for keyword in required_keywords if keyword not in full_text]
    if missing:
        raise RuntimeError(f"说明书自检失败，缺少关键词：{', '.join(missing)}")

    heading_count = sum(1 for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading"))
    figure_count = len(document.inline_shapes)
    if heading_count < 30:
        raise RuntimeError(f"说明书自检失败，标题数量异常：{heading_count}")
    if figure_count < 18:
        raise RuntimeError(f"说明书自检失败，图片数量异常：{figure_count}")
    return {"headings": heading_count, "figures": figure_count, "paragraphs": len(document.paragraphs)}


def parse_args() -> argparse.Namespace:
    """解析命令行参数。"""

    parser = argparse.ArgumentParser(description="生成通用CAE分析平台软件使用说明书")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="输出 docx 路径")
    parser.add_argument("--title", default=DEFAULT_TITLE, help="软件标题")
    parser.add_argument("--version", default=DEFAULT_VERSION, help="软件版本号")
    parser.add_argument("--skip-backup", action="store_true", help="覆盖前不生成备份")
    return parser.parse_args()


def main() -> int:
    """脚本入口。"""

    args = parse_args()
    output_path: Path = args.output
    backup_path = None if args.skip_backup else backup_existing_file(output_path)
    build_manual(output_path=output_path, title=args.title, version=args.version)
    summary = inspect_manual(output_path)
    print(f"说明书已生成：{output_path}")
    if backup_path is not None:
        print(f"旧文档备份：{backup_path}")
    print(
        "自检通过："
        f"标题 {summary['headings']} 个，"
        f"图片 {summary['figures']} 张，"
        f"段落 {summary['paragraphs']} 段。"
    )
    print(f"占位图目录：{PLACEHOLDER_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
